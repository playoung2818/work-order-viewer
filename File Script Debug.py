import os
import logging
from flask import Flask, jsonify
from flask_sqlalchemy import SQLAlchemy
from docx import Document
import pdfplumber # type: ignore
import pandas as pd
import json

# Logging Configuration
logging.basicConfig(level=logging.INFO)

# Flask and Database Configuration
app = Flask(__name__)
app.config['SQLALCHEMY_DATABASE_URI'] = 'postgresql://postgres:Czheyuan0227%40@localhost:5432/File_Log'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

db = SQLAlchemy(app)

# Database Models
class PDFFileLog(db.Model):
    __tablename__ = 'pdf_file_log'
    id = db.Column(db.Integer, primary_key=True)
    order_id = db.Column(db.String(255), nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    file_path = db.Column(db.String(255), nullable=False)
    extracted_data = db.Column(db.JSON, nullable=True)
    __table_args__ = (db.UniqueConstraint('order_id', 'file_name', name='unique_pdf_entry'),)

class WordFileLog(db.Model):
    __tablename__ = 'word_file_log'
    id = db.Column(db.Integer, primary_key=True)
    order_id = db.Column(db.String(255), nullable=False)
    file_name = db.Column(db.String(255), nullable=False)
    product_details = db.Column(db.JSON, nullable=False)
    file_path = db.Column(db.String(255), nullable=False)
    status = db.Column(db.String(50), default="Not Picked")

    __table_args__ = (db.UniqueConstraint('order_id', 'file_name', name='unique_word_entry'),)

# Helper Functions
def validate_paths(paths):
    for path in paths:
        if not os.path.exists(path):
            logging.error(f"Path does not exist: {path}")
        else:
            logging.info(f"Valid path: {path}")

# Function to Extract Tables from PDF
def extract_pdf_tables(pdf_path):
    tables_list = []
    
    with pdfplumber.open(pdf_path) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            tables = page.extract_table()
            
            # Check if the table extraction is successful
            if not tables or len(tables) < 2:
                logging.warning(f"No valid tables found on page {page_num} of {pdf_path}")
                continue  # Skip to the next page
            
            df = pd.DataFrame(tables[1:], columns=tables[0])  # First row as headers
            
            # Ensure expected columns exist
            expected_columns = {'Item', 'Ordered', 'Description'}
            missing_columns = expected_columns - set(df.columns)
            if missing_columns:
                logging.warning(f"Missing expected columns {missing_columns} in {pdf_path}, page {page_num}")
                continue  # Skip this page if required columns are missing

            df = df[df['Description'].notna() & (df['Description'].str.strip() != "")]
            df.insert(0, "Page", page_num)
            tables_list.append(df)

    if tables_list:
        final_df = pd.concat(tables_list, ignore_index=True)
        final_df.drop_duplicates(subset=['Item', 'Ordered'], inplace=True)  
        final_df = final_df[final_df['Item'] != 'Forwarding Charge']
        
        json_data = final_df[['Item', 'Ordered']].to_dict(orient="records")
        return json.dumps(json_data)  
    return json.dumps([])  


def process_pdf_files(folder_path):
    pdf_data = {}
    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.pdf'):
                file_path = os.path.join(root, file)
                order_id = os.path.splitext(file)[0]  
                logging.info(f"Processing PDF file: {file}")

                # Extract table data from PDF
                extracted_data = extract_pdf_tables(file_path)

                # ✅ Ensure JSON parsing
                try:
                    extracted_data_list = json.loads(extracted_data)
                except json.JSONDecodeError as e:
                    logging.error(f"Error parsing extracted_data: {e}")
                    extracted_data_list = []

                extracted_data = "\n".join(f"{entry['Item']}\t{entry['Ordered']}" 
                                           for entry in extracted_data_list 
                                           if entry.get('Item') and entry.get('Ordered'))

                # Store file details
                pdf_data[order_id] = {
                    "file_name": file,
                    "file_path": file_path,
                    "extracted_data": extracted_data
                }
    return pdf_data

def process_word_files(folder_path):
    word_data = {}
    if not os.path.exists(folder_path):
        logging.error(f"Folder path does not exist: {folder_path}")
        return word_data

    for root, _, files in os.walk(folder_path):
        for file in files:
            if file.lower().endswith('.docx'):
                try:
                    file_path = os.path.join(root, file)
                    order_id = os.path.splitext(file)[0]
                    logging.info(f"Processing Word file: {file}")
                    product_details = extract_product_details_from_word(file_path)
                    word_data[order_id] = {
                        "file_name": file,
                        "file_path": file_path,
                        "product_details": product_details
                    }
                except Exception as e:
                    logging.error(f"Error processing Word file {file}: {e}")
    return word_data

def extract_product_details_from_word(file_path):
    try:
        if not os.path.exists(file_path):
            logging.warning(f"Word file path does not exist: {file_path}")
            return []

        document = Document(file_path)
        if not document.tables:
            logging.warning(f"No tables found in Word file: {file_path}")
            return []

        table = document.tables[0]
        product_details = []
        for i, row in enumerate(table.rows[1:]):  # Skip header row
            cells = row.cells
            if len(cells) < 4:
                logging.warning(f"Row {i + 1} in {file_path} has insufficient cells.")
                continue

            product_details.append({
                "product_number": cells[0].text.strip(),
                "qty": cells[1].text.strip(),
                "sn": cells[2].text.strip(),
                "notes": cells[3].text.strip()
            })
        return product_details
    except Exception as e:
        logging.error(f"Error processing Word file {file_path}: {e}")
        return []

# Function to Load PDF Data into Database
def load_pdf_files_to_db(pdf_data):
    for order_id, data in pdf_data.items():
        try:
            # Check if the entry already exists
            existing_entry = PDFFileLog.query.filter_by(order_id=order_id, file_name=data['file_name']).first()
            if existing_entry:
                logging.info(f"Duplicate PDF entry detected for Order ID {order_id}, File: {data['file_name']}. Skipping insert.")
                continue  # Skip this entry
            
            logging.info(f"Inserting into DB: Order ID: {order_id}, File Name: {data['file_name']}")
            
            db.session.add(PDFFileLog(
                order_id=order_id,
                file_name=data['file_name'],
                file_path=data['file_path'],
                extracted_data=json.dumps(data['extracted_data'])  # ✅ Store JSON safely
            ))
        except Exception as e:
            db.session.rollback()
            logging.error(f"Database insert error for order ID {order_id}: {e}")

    db.session.commit()
    logging.info("All PDF file data loaded into the database.")


def load_word_files_to_db(word_data):
    for order_id, data in word_data.items():
        try:
            existing_entry = WordFileLog.query.filter_by(order_id=order_id, file_name=data['file_name']).first()
            if existing_entry:
                logging.info(f"Duplicate Word entry detected for Order ID {order_id}. Skipping.")
                continue

            db.session.add(WordFileLog(
                order_id=order_id,
                file_name=data['file_name'],
                product_details=data['product_details'],
                file_path=data['file_path']
            ))
        except Exception as e:
            logging.error(f"Error saving Word file data for order ID {order_id}: {e}")
    db.session.commit()
    logging.info("Word files successfully saved to database.")

# Main Functions
def process_all_work_order_pdfs():
    WO_PDF_FOLDER = r"\\NEOUSYSSERVER\Drive D\QuickBooks\2- Year 2024\Work Order- WO"
    WO_PDF_FOLDER2 = r"\\NEOUSYSSERVER\Drive D\QuickBooks\3- Year 2025\Work Order- WO"
    logging.info("Processing all Work Order PDF files...")
    pdf_data = process_pdf_files(WO_PDF_FOLDER)
    pdf_data.update(process_pdf_files(WO_PDF_FOLDER2))
    if pdf_data:
        load_pdf_files_to_db(pdf_data)
    else:
        logging.info("No PDF files found in the Work Order folder.")
    logging.info("All Work Order PDF files processed and saved.")

def process_all_work_order_words():
    WO_WORD_FOLDER = r"C:\Users\Admin\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2024"
    WO_WORD_FOLDER2 = r"C:\Users\Admin\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2025"
    logging.info("Processing all Work Order Word files...")
    word_data = process_word_files(WO_WORD_FOLDER)
    word_data.update(process_word_files(WO_WORD_FOLDER2))
    if word_data:
        load_word_files_to_db(word_data)
    else:
        logging.info("No Word files found in the Work Order folder.")
    logging.info("All Work Order Word files processed and saved.")

# ✅ API Endpoint (was from API_server.py)
@app.route('/api/word-files', methods=['GET'])
def get_all_word_files():
    try:
        word_files = WordFileLog.query.all()
        file_list = [
            {"order_id": file.order_id, "file_name": file.file_name, "status": "Picked"}
            for file in word_files
        ]
        for file in word_files:
            file.status = "Picked"
        db.session.commit()
        return jsonify({"word_files": file_list}), 200
    except Exception as e:
        logging.error(f"Error fetching Word files: {e}")
        return jsonify({"error": str(e)}), 500

# Flask Main Context
if __name__ == "__main__":
    with app.app_context():
        logging.info("Initializing database tables...")
        try:
            db.create_all()
            logging.info("Database initialized successfully.")
        except Exception as e:
            logging.error(f"Database initialization error: {e}")
            exit(1)

        # Validate paths
        validate_paths([
            r"\\NEOUSYSSERVER\Drive D\QuickBooks\2- Year 2024\Work Order- WO",
            r"\\NEOUSYSSERVER\Drive D\QuickBooks\3- Year 2025\Work Order- WO",
            r"C:\Users\Admin\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2024",
            r"C:\Users\Admin\OneDrive - neousys-tech\Share NTA Warehouse\02 Work Order- Word file\Work Order 2025"
        ])

        # Process files
        process_all_work_order_pdfs()
        process_all_work_order_words()

    logging.info("🚀 Starting API server on http://localhost:5001")
    app.run(host="0.0.0.0", port=5001, debug=True)


